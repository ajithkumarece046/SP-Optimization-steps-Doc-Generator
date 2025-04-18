import streamlit as st
import re
import pandas as pd
import os
import json
from io import StringIO, BytesIO
# from dotenv import load_dotenv # Removed as using secrets
from openai import AzureOpenAI
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_STYLE_TYPE # Added WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import traceback # Added for better error reporting

# Set page configuration
st.set_page_config(
    page_title="SQL Stored Procedure Analyzer",
    page_icon="🧰",
    layout="wide"
)

# --- Function to create Word document ---
def create_word_document(analysis):
    """Creates a Word document (.docx) summarizing the analysis."""
    try:
        # Create a new Document
        doc = Document()

        # --- Document Styles (Optional but recommended) ---
        # Define styles if needed, e.g., for code blocks
        styles = doc.styles
        # try:
        #     code_style = styles['Code'] # Check if 'Code' style exists
        # except KeyError:
        #     # Create a base style for code if 'Code' doesn't exist
        #     from docx.shared import RGBColor
        #     style = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH) # Use unique name
        #     style.font.name = 'Courier New'
        #     style.font.size = Pt(10)
        #     # Optional: Add light background shading
        #     # style.paragraph_format.shading.background_color = RGBColor(0xF2, 0xF2, 0xF2)
        #     # Optional: Add indentation
        #     style.paragraph_format.left_indent = Inches(0.25)

        # --- Title ---
        title = doc.add_heading('SQL Stored Procedure Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Procedure Name ---
        doc.add_heading('Stored Procedure Name:', level=1)
        proc_name_para = doc.add_paragraph()
        proc_name_para.add_run(analysis.get('procedure_name', 'N/A')).bold = True # Use .get for safety

        # --- Scope ---
        doc.add_heading('Scope:', level=1)
        doc.add_paragraph(analysis.get('scope', 'N/A')) # Use .get for safety

        # --- Optimization Steps ---
        doc.add_heading('Optimization Steps:', level=1)

        optimizations = analysis.get("optimizations", []) # Use .get for safety
        if not optimizations:
            doc.add_paragraph("No optimization suggestions were generated.")
        else:
            for i, opt in enumerate(optimizations, 1):
                # Step heading
                opt_type = opt.get("type", "N/A")
                doc.add_heading(f'Step {i}: {opt_type}', level=2)

                # Existing Logic
                doc.add_heading('Existing Logic:', level=3)
                existing_logic = opt.get("existing_logic", "")
                if existing_logic:
                    existing_code_para = doc.add_paragraph()
                    run = existing_code_para.add_run(existing_logic)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    existing_code_para.paragraph_format.left_indent = Inches(0.25)
                    # existing_code_para.style = 'CodeStyle' # Apply style if defined
                else:
                    doc.add_paragraph("N/A")


                # Optimized Logic
                doc.add_heading('Optimized Logic:', level=3)
                optimized_logic = opt.get("optimized_logic", "")
                if optimized_logic:
                    optimized_code_para = doc.add_paragraph()
                    run = optimized_code_para.add_run(optimized_logic)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    optimized_code_para.paragraph_format.left_indent = Inches(0.25)
                    # optimized_code_para.style = 'CodeStyle' # Apply style if defined
                else:
                    doc.add_paragraph("N/A")


                # Explanation
                explanation_para = doc.add_paragraph()
                explanation_text = explanation_para.add_run(opt.get("explanation", "N/A"))
                explanation_text.italic = True

                # Add separator paragraph
                separator = doc.add_paragraph()
                separator.add_run('_' * 50) # Slightly longer separator

        # --- Summary Table ---
        doc.add_heading('Summary:', level=1)

        # Create table data
        table_data = []
        if optimizations: # Only build table if there are optimizations
            for opt in optimizations:
                table_data.append({
                    "Type of Change": opt.get("type", "N/A"),
                    "Line Number": str(opt.get("line_number", "N/A")), # Ensure string type for line number
                    "Original Code Snippet": opt.get("existing_logic", ""), # <-- Display full snippet
                    "Optimized Code Snippet": opt.get("optimized_logic", ""), # <-- Display full snippet
                    "Optimization Explanation": opt.get("explanation", "")
                })

        # Add table to document only if there's data
        if table_data:
            num_rows = 1 + len(table_data)
            num_cols = 5
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid' # Use a built-in style
            table.autofit = False # Disable autofit to respect column widths better
            table.allow_autofit = False # Ensure autofit is off

            # --- Table Header ---
            header_cells = table.rows[0].cells
            headers = ['Type of Change', 'Line Number', 'Original Code Snippet', 'Optimized Code Snippet', 'Optimization Explanation']
            for i, header_text in enumerate(headers):
                cell = header_cells[i]
                cell.text = '' # Clear existing content
                p = cell.paragraphs[0]
                run = p.add_run(header_text)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center align header

                # Optional: Add shading to header row
                tcPr = cell._tc.get_or_add_tcPr()
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), "D9D9D9") # Slightly darker gray for header
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:color'), 'auto')
                tcPr.append(shading_elm)


            # --- Table Data Rows ---
            for i, item in enumerate(table_data):
                row_cells = table.rows[i+1].cells
                row_cells[0].text = item['Type of Change']
                row_cells[1].text = item['Line Number']
                row_cells[2].text = item['Original Code Snippet'] # Full code here
                row_cells[3].text = item['Optimized Code Snippet'] # Full code here
                row_cells[4].text = item['Optimization Explanation']

                # Optional: Apply code font to code snippet cells
                for cell_idx in [2, 3]:
                    cell = row_cells[cell_idx]
                    for paragraph in cell.paragraphs:
                        # Ensure cell vertical alignment is top for code readability
                        cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.TOP
                        for run in paragraph.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9) # Slightly smaller font for table code

                # Apply alternate row shading
                if (i + 1) % 2 == 0:  # Even data rows (index 1, 3, ... which are rows 2, 4, ...) get shading
                    for cell in row_cells:
                        tcPr = cell._tc.get_or_add_tcPr()
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), "F2F2F2") # Light gray color
                        shading_elm.set(qn('w:val'), 'clear')
                        shading_elm.set(qn('w:color'), 'auto')
                        tcPr.append(shading_elm)

            # --- Set Table Column Widths (Adjusted) ---
            # These widths are suggestions; Word's layout engine might adjust them.
            try:
                # Ensure table layout allows fixed widths
                tbl_pr = table._tbl.tblPr
                tbl_layout = OxmlElement('w:tblLayout')
                tbl_layout.set(qn('w:type'), 'fixed')
                tbl_pr.append(tbl_layout)

                table.columns[0].width = Inches(1.0) # Type of Change
                table.columns[1].width = Inches(0.5) # Line Number
                table.columns[2].width = Inches(2.2) # Original Code Snippet (Increased)
                table.columns[3].width = Inches(2.2) # Optimized Code Snippet (Increased)
                table.columns[4].width = Inches(1.6) # Explanation (Adjusted)
                # Total Width ~ 7.5 inches (might exceed standard page width with margins, word will wrap)
            except IndexError:
                 st.warning("Could not set all table column widths. Check number of columns.")
            except Exception as e:
                 st.warning(f"Error setting column widths: {e}")

        else: # Handles case where analysis had optimizations=None or optimizations=[]
            if not optimizations: # If optimizations was explicitly checked and found empty earlier
                 pass # Already handled under "Optimization Steps" section
            else: # If table_data is empty for other reasons (e.g., filtering)
                 doc.add_paragraph("Summary table could not be generated.")


        # --- Save the document to a BytesIO object ---
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io

    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        st.sidebar.expander("Word Gen Error Details", expanded=False).code(traceback.format_exc())
        return None # Return None if document generation fails

# --- Function to analyze stored procedure using Azure OpenAI ---
def analyze_stored_procedure(file_content):
    """Analyzes SQL content using Azure OpenAI and returns parsed JSON."""
    if not file_content or not file_content.strip():
         st.warning("Cannot analyze empty SQL content.")
         return None
    try:
        # Load credentials securely from Streamlit secrets
        azure_openai_endpoint = st.secrets["AZURE_OPENAI_ENDPOINT"]
        azure_openai_key = st.secrets["AZURE_OPENAI_API_KEY"]
        azure_api_version = st.secrets["API_VERSION"]
        deployment_name = st.secrets.get("DEPLOYMENT_NAME", "gpt-4o-mini") # Use provided name or default

        # Validate credentials
        if not all([azure_openai_endpoint, azure_openai_key, azure_api_version]):
            st.error("Missing required secrets (AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_KEY, API_VERSION). Please check your Streamlit secrets configuration.")
            return None

        # Initialize Azure OpenAI client
        client = AzureOpenAI(
            api_key=azure_openai_key,
            api_version=azure_api_version,
            azure_endpoint=azure_openai_endpoint
        )

        prompt = f"""
        Analyze the following SQL stored procedure and return your analysis ONLY in JSON format:

        ```sql
        {file_content}
        ```

        **Instructions:**
        1.  Identify the stored procedure name (value for "procedure_name").
        2.  Describe the scope/purpose of the stored procedure in 4-5 concise lines (value for "scope").
        3.  Identify up to 5 high-priority optimization opportunities. Focus on significant performance impacts like:
            *   Replacing cursors with set-based operations (e.g., CTEs, derived tables).
            *   Consolidating multiple UPDATE/DELETE statements targeting the same rows.
            *   Optimizing or adding necessary indexes, especially for JOINs and WHERE clauses (suggest index creation statements if applicable).
            *   Identifying and rewriting inefficient query patterns (e.g., correlated subqueries, functions in WHERE clauses).
            *   Detecting unused variables or temporary tables.
            *   Simplifying complex logic where possible.
            *   Parameter sniffing issues if identifiable.
        4.  For each optimization opportunity, create a JSON object within the "optimizations" array containing:
            *   `type`: A short description (e.g., "Replace Cursor", "Combine Updates", "Add Index").
            *   `line_number`: The approximate starting line number or range (e.g., "55" or "55-60") where the existing logic is found. Use "N/A" if not applicable.
            *   `existing_logic`: The *complete*, relevant block of the original SQL code that needs modification. Include enough context. Ensure this is a single string, escaping newlines if necessary within the JSON string.
            *   `optimized_logic`: The *complete*, suggested replacement SQL code, including any necessary surrounding syntax. Ensure this is a single string, escaping newlines if necessary within the JSON string.
            *   `explanation`: A brief explanation of *why* the change is beneficial (e.g., "Reduces loops, improves set-based processing", "Minimizes I/O by combining DML operations", "Speeds up lookups on the temp table").
        5.  Provide a brief overall summary in the "summary" object containing:
            *   `original_performance_issues`: Key performance problems identified (string).
            *   `optimization_impact`: Expected overall impact, e.g., "Significant performance improvement expected", "Moderate reduction in execution time" (string).
            *   `implementation_difficulty`: Estimated effort, e.g., "Low", "Medium", "High" (string).

        **Output Format (Strict JSON):**
        Your entire response MUST be a single, valid JSON object adhering precisely to this schema. Do NOT include any text, explanations, apologies, or markdown formatting (like ```json) before or after the JSON object.

        ```json
        {{
          "procedure_name": "string",
          "scope": "string",
          "optimizations": [
            {{
              "type": "string",
              "line_number": "string",
              "existing_logic": "string (full SQL code snippet as a single JSON string)",
              "optimized_logic": "string (full SQL code snippet as a single JSON string)",
              "explanation": "string"
            }}
          ],
          "summary": {{
            "original_performance_issues": "string",
            "optimization_impact": "string",
            "implementation_difficulty": "string"
          }}
        }}
        ```
        """

        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": "You are an expert SQL database optimizer. Your response MUST be a single, valid JSON object matching the requested schema, with no surrounding text or markdown."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2, # Lower temperature for more deterministic JSON output
            max_tokens=4000, # Increased max tokens for potentially long code snippets
            response_format={"type": "json_object"} # Explicitly request JSON response
        )

        # Extract the JSON content
        analysis_result_str = response.choices[0].message.content

        # Debug: Display raw response
        # st.sidebar.expander("Debug Raw LLM Response", expanded=False).code(analysis_result_str)

        # Attempt to parse the JSON
        try:
            analysis_data = json.loads(analysis_result_str)
            # Basic validation (check if required keys exist)
            required_keys = ["procedure_name", "scope", "optimizations", "summary"]
            if not all(k in analysis_data for k in required_keys):
                 st.warning("Warning: The analysis response might be missing some expected top-level keys (procedure_name, scope, optimizations, summary).")
            if "optimizations" in analysis_data and isinstance(analysis_data["optimizations"], list):
                opt_keys = ["type", "line_number", "existing_logic", "optimized_logic", "explanation"]
                for i, opt in enumerate(analysis_data["optimizations"]):
                     if not all(k in opt for k in opt_keys):
                          missing = [k for k in opt_keys if k not in opt]
                          st.warning(f"Warning: Optimization step {i+1} might be missing required keys: {', '.join(missing)}")
            elif "optimizations" not in analysis_data:
                 # If optimizations key is missing, add an empty list for downstream safety
                 analysis_data["optimizations"] = []


            # Ensure summary exists, even if empty
            if "summary" not in analysis_data:
                 analysis_data["summary"] = {
                      "original_performance_issues": "N/A",
                      "optimization_impact": "N/A",
                      "implementation_difficulty": "N/A"
                 }


            return analysis_data

        except json.JSONDecodeError as e:
            st.error(f"Failed to parse the JSON response from the AI model: {str(e)}")
            st.error("The received response was:")
            st.code(analysis_result_str) # Show the problematic response
            return None
        except Exception as e: # Catch other potential errors during parsing/validation
             st.error(f"An unexpected error occurred while processing the analysis response: {str(e)}")
             st.code(analysis_result_str)
             return None

    except Exception as e:
        st.error(f"Error during analysis API call: {str(e)}")
        st.sidebar.expander("API Error Details", expanded=False).code(traceback.format_exc())
        return None

# --- UI Components ---
st.title("🚀 SQL Stored Procedure Analyzer")
st.write("Upload a SQL stored procedure file (.sql) or use the sample code for AI-powered optimization analysis and report generation.")

st.sidebar.header("Configuration Info")
st.sidebar.info(f"""
Uses Azure OpenAI via Streamlit Secrets:
- `AZURE_OPENAI_ENDPOINT`
- `AZURE_OPENAI_API_KEY`
- `API_VERSION`
- `DEPLOYMENT_NAME` (optional, defaults to 'gpt-4o-mini')
""")

# Check if secrets are loaded
if not all(s in st.secrets for s in ["AZURE_OPENAI_ENDPOINT", "AZURE_OPENAI_API_KEY", "API_VERSION"]):
    st.error("🛑 Critical secrets missing! Please configure `AZURE_OPENAI_ENDPOINT`, `AZURE_OPENAI_API_KEY`, and `API_VERSION` in your Streamlit secrets.")
    st.stop() # Halt execution if secrets aren't available

st.markdown("---")

# --- Main Area ---

col1, col2 = st.columns([1, 1]) # Create two columns for layout

# ========================= COLUMN 1: INPUT =========================
with col1:
    st.subheader("1. Provide SQL Code")

    # --- File Upload Handling (using multi-stage approach) ---
    # Phase 1: Detect upload, store object, mark as seen, and trigger rerun
    uploaded_file_obj = st.file_uploader("Upload SQL Stored Procedure File", type=["sql"], label_visibility="collapsed", key="sql_uploader")

    if uploaded_file_obj is not None and uploaded_file_obj != st.session_state.get('processed_uploaded_file_obj'):
        # Store the new file object to be processed on the next run
        st.session_state['newly_uploaded_file_obj'] = uploaded_file_obj
        # Mark this specific file object as "seen" to prevent re-processing if the user interacts elsewhere
        st.session_state['processed_uploaded_file_obj'] = uploaded_file_obj
        # Clear sample state if a file is uploaded
        if 'sample_sql_loaded' in st.session_state:
            del st.session_state['sample_sql_loaded']
        st.rerun() # Trigger immediate rerun to handle the uploaded file

    # Phase 2: Process the stored file object on the script run *after* upload detection
    if 'newly_uploaded_file_obj' in st.session_state and st.session_state['newly_uploaded_file_obj'] is not None:
        uploaded_file_to_process = st.session_state['newly_uploaded_file_obj']
        try:
            # Decode and update the session state bound to the text area
            sql_content_from_file = uploaded_file_to_process.getvalue().decode("utf-8")
            st.session_state['sql_input'] = sql_content_from_file
            st.session_state['file_name'] = os.path.splitext(uploaded_file_to_process.name)[0]
            # Important: Clear the trigger variable *after* successful processing
            st.session_state['newly_uploaded_file_obj'] = None
            # Provide feedback (optional)
            # st.success(f"Loaded content from: {uploaded_file_to_process.name}")
        except UnicodeDecodeError:
             st.error(f"Error decoding file '{uploaded_file_to_process.name}'. Please ensure it is UTF-8 encoded.")
             st.session_state['newly_uploaded_file_obj'] = None # Clear trigger
        except Exception as e:
            st.error(f"Error reading uploaded file: {e}")
            # Important: Clear the trigger variable even on error
            st.session_state['newly_uploaded_file_obj'] = None
        # No rerun needed here, the text_area below will pick up the updated 'sql_input' state

    # --- Sample SQL Button ---
    if st.button("Load Sample SQL"):
        sample_sql = """
CREATE PROCEDURE usp_GetCustomerOrders_Inefficient
    @CustomerId INT,
    @OrderDateThreshold DATETIME = NULL
AS
BEGIN
    SET NOCOUNT ON;

    -- Check if threshold is provided, default if not
    IF @OrderDateThreshold IS NULL
        SET @OrderDateThreshold = '1900-01-01'; -- Default to a very old date

    -- Temp table for intermediate results
    CREATE TABLE #CustomerOrders (
        OrderId INT PRIMARY KEY,
        OrderDate DATETIME,
        TotalAmount DECIMAL(10, 2),
        ItemCount INT
    );

    -- Insert initial orders (potential performance issue if Orders table is large)
    INSERT INTO #CustomerOrders (OrderId, OrderDate, TotalAmount)
    SELECT OrderId, OrderDate, OrderValue
    FROM Orders
    WHERE CustomerId = @CustomerId AND OrderDate >= @OrderDateThreshold;

    -- Cursor to update item count for each order - VERY INEFFICIENT
    DECLARE @CurrentOrderId INT;
    DECLARE order_cursor CURSOR LOCAL FAST_FORWARD FOR
        SELECT OrderId FROM #CustomerOrders;

    OPEN order_cursor;
    FETCH NEXT FROM order_cursor INTO @CurrentOrderId;

    WHILE @@FETCH_STATUS = 0
    BEGIN
        -- Update item count using a correlated subquery inside the loop
        UPDATE #CustomerOrders
        SET ItemCount = (SELECT COUNT(*) FROM OrderDetails WHERE OrderId = @CurrentOrderId)
        WHERE OrderId = @CurrentOrderId;

        -- Simulate some processing delay or complex logic
        WAITFOR DELAY '00:00:00.010'; -- Artificial delay

        -- Update Order status individually - Inefficient
        UPDATE Orders SET Status = 'Processing Started' WHERE OrderId = @CurrentOrderId;
        UPDATE Orders SET LastModifiedDate = GETDATE() WHERE OrderId = @CurrentOrderId;

        FETCH NEXT FROM order_cursor INTO @CurrentOrderId;
    END;

    CLOSE order_cursor;
    DEALLOCATE order_cursor;

    -- Final SELECT - Join back to Customers (maybe unnecessary if only order data needed)
    SELECT
        c.CustomerName,
        co.OrderId,
        co.OrderDate,
        co.TotalAmount,
        co.ItemCount
    FROM
        #CustomerOrders co
    INNER JOIN Customers c ON c.CustomerId = @CustomerId -- Assuming @CustomerId matches the customer for these orders
    ORDER BY
        co.OrderDate DESC;

    -- Cleanup
    DROP TABLE #CustomerOrders;

END;
GO
        """
        # Directly update state - safe within button callback context
        st.session_state['sql_input'] = sample_sql
        st.session_state['file_name'] = "sample_procedure"
        st.session_state['sample_sql_loaded'] = True # Flag that sample is loaded
        # Clear file upload tracking state if sample is loaded
        st.session_state['processed_uploaded_file_obj'] = None
        st.session_state['newly_uploaded_file_obj'] = None
        st.success("Sample SQL loaded!")
        st.rerun() # Rerun to ensure UI consistency

    # --- Text Area ---
    # Initialize 'sql_input' in session_state if it doesn't exist
    if 'sql_input' not in st.session_state:
        st.session_state['sql_input'] = ""

    # The text area's value is now primarily controlled by session_state['sql_input']
    # which is updated either by file processing (Phase 2) or the sample button.
    sql_input_from_textarea = st.text_area("Or paste SQL code here:", height=300, key="sql_input")

    # --- Determine final content and filename for analysis ---
    # The text area state IS the primary source now.
    sql_content = st.session_state.get('sql_input', '')

    # Determine the filename based on the source
    if st.session_state.get('sample_sql_loaded'):
        file_name = "sample_procedure"
    elif st.session_state.get('processed_uploaded_file_obj') and st.session_state.get('file_name'):
        # If an upload was processed, keep its name even if text area is edited
        file_name = st.session_state.get('file_name')
    elif sql_content:
        # If content exists (pasted/edited) and no upload/sample context, use default
        file_name = "pasted_procedure"
    else:
        # Default if everything is empty
        file_name = "analysis_report"

    # Store the determined filename back to state for consistent use
    st.session_state['current_file_name_base'] = file_name


# ========================= COLUMN 2: ANALYSIS & DOWNLOAD =========================
with col2:
    st.subheader("2. Analyze and Download")
    # Button is enabled only if there is content in the text area state
    analyze_button = st.button(
        "Analyze SQL Procedure",
        type="primary",
        disabled=not st.session_state.get('sql_input', '').strip() # Check if sql_input state is non-empty
    )

    # Store analysis results in session state to persist across interactions
    if analyze_button:
        # Run analysis using the content currently in session state
        current_sql_content = st.session_state.get('sql_input', '')
        if current_sql_content.strip():
             with st.spinner("🤖 Analyzing stored procedure... This may take a minute."):
                st.session_state.analysis = analyze_stored_procedure(current_sql_content)
        else:
             st.warning("Cannot analyze empty SQL content.")
             # Clear previous analysis if input is now empty
             if 'analysis' in st.session_state:
                  del st.session_state['analysis']


    # Display results if analysis is available in session state
    if 'analysis' in st.session_state and st.session_state.analysis:
        analysis = st.session_state.analysis
        st.success("Analysis Complete!")

        # Display procedure name and summary
        st.markdown(f"#### 🏷️ Procedure Name: `{analysis.get('procedure_name', 'N/A')}`")
        with st.expander("View Scope & Summary", expanded=False):
             st.markdown("**Scope:**")
             st.write(analysis.get('scope', 'N/A'))
             st.markdown("**Analysis Summary:**")
             summary = analysis.get('summary', {}) # Default to empty dict
             st.write(f"- **Identified Issues:** {summary.get('original_performance_issues', 'N/A')}")
             st.write(f"- **Expected Impact:** {summary.get('optimization_impact', 'N/A')}")
             st.write(f"- **Implementation Difficulty:** {summary.get('implementation_difficulty', 'N/A')}")

        # Display optimization steps
        st.markdown("#### ✨ Optimization Suggestions:")
        optimizations = analysis.get("optimizations", []) # Default to empty list

        if not optimizations:
             st.info("✅ No specific optimization suggestions were provided in the analysis.")
        else:
            for i, opt in enumerate(optimizations, 1):
                opt_type = opt.get('type', 'N/A')
                opt_line = opt.get('line_number', 'N/A')
                expander_title = f"Step {i}: {opt_type} (Line: {opt_line})"
                with st.expander(expander_title, expanded=(i==1)): # Expand first step
                    col_exist, col_optim = st.columns(2)
                    with col_exist:
                        st.markdown("**Existing Logic:**")
                        st.code(opt.get("existing_logic", "N/A"), language="sql")
                    with col_optim:
                        st.markdown("**Optimized Logic:**")
                        st.code(opt.get("optimized_logic", "N/A"), language="sql")

                    st.markdown("**Explanation:**")
                    st.caption(f"> {opt.get('explanation', 'N/A')}")

        # --- Download Section ---
        st.markdown("---")
        st.subheader("📥 Download Report")

        # Use the filename base determined in col1
        report_filename_base = st.session_state.get('current_file_name_base', 'sql_analysis')

        # Generate Word doc bytes
        docx_bytes = None
        try:
            with st.spinner("Generating Word document..."):
                docx_bytes = create_word_document(analysis)
        except Exception as e:
            st.error(f"Failed to generate Word document bytes: {e}") # Log error


        if docx_bytes:
            st.download_button(
                label="⬇️ Download as Word (.docx)",
                data=docx_bytes,
                file_name=f"{report_filename_base}_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key='docx-download'
            )
        else:
             st.warning("Word document could not be generated.")

        # --- Markdown Download ---
        try:
            # Create summary DataFrame for Markdown export
            summary_data_md = []
            for opt in analysis.get("optimizations", []):
                summary_data_md.append({
                    "Type": opt.get("type", "N/A"),
                    "Line": opt.get("line_number", "N/A"),
                    "Original Code": f"```sql\n{opt.get('existing_logic', '')}\n```", # Format as code block
                    "Optimized Code": f"```sql\n{opt.get('optimized_logic', '')}\n```", # Format as code block
                    "Explanation": opt.get("explanation", "")
                })
            summary_df_md = pd.DataFrame(summary_data_md)

            # Generate Markdown content
            report_md = f"# SQL Stored Procedure Analysis Report\n\n"
            report_md += f"## Procedure Name: `{analysis.get('procedure_name', 'N/A')}`\n\n"
            report_md += f"## Scope:\n{analysis.get('scope', 'N/A')}\n\n"
            report_md += f"## Analysis Summary:\n"
            summary = analysis.get('summary', {})
            report_md += f"- **Identified Issues:** {summary.get('original_performance_issues', 'N/A')}\n"
            report_md += f"- **Expected Impact:** {summary.get('optimization_impact', 'N/A')}\n"
            report_md += f"- **Implementation Difficulty:** {summary.get('implementation_difficulty', 'N/A')}\n\n"

            report_md += "## Optimization Steps:\n"
            optimizations_md = analysis.get("optimizations", [])
            if not optimizations_md:
                 report_md += "No specific optimization suggestions were provided.\n"
            else:
                for i, opt in enumerate(optimizations_md, 1):
                    opt_type = opt.get('type', 'N/A')
                    opt_line = opt.get('line_number', 'N/A')
                    report_md += f"\n### Step {i}: {opt_type} (Line: {opt_line})\n\n"
                    report_md += f"**Existing Logic:**\n```sql\n{opt.get('existing_logic', 'N/A')}\n```\n\n"
                    report_md += f"**Optimized Logic:**\n```sql\n{opt.get('optimized_logic', 'N/A')}\n```\n\n"
                    report_md += f"**Explanation:**\n> {opt.get('explanation', 'N/A')}\n\n---\n"

            if not summary_df_md.empty:
                report_md += "\n## Summary Table:\n\n"
                report_md += summary_df_md.to_markdown(index=False)

            st.download_button(
                label="⬇️ Download as Markdown (.md)",
                data=report_md.encode('utf-8'), # Encode to bytes
                file_name=f"{report_filename_base}_analysis.md",
                mime="text/markdown",
                key='md-download'
            )
        except Exception as e:
             st.warning(f"Could not generate Markdown download: {e}")


    elif analyze_button and 'analysis' in st.session_state and not st.session_state.analysis:
        # Handle case where analysis button was clicked but analysis failed or returned None
        st.error("Analysis could not be completed or returned no results. Please check the SQL content and any error messages above or in the sidebar.")

    elif not st.session_state.get('sql_input','').strip(): # If the text area state is empty
         st.info("Upload a .sql file or paste/load SQL code in the text area on the left to enable analysis.")


# --- Footer ---
st.markdown("---")
st.caption("SQL Stored Procedure Analyzer powered by Azure OpenAI")